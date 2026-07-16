"""
DOCX Exporter — 学霸笔记 .docx 生成器

Minimal, single-function module. No external deps beyond python-docx.
Fonts use CJK fallback chain: 黑体/宋体 with SimHei/SimSun alternatives.
"""

import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---- Helper: set CJK font on a run ----
def _set_run_font(run, font_name_cn="宋体", font_name_en="Calibri", size_pt=11,
                  bold=False, color=None):
    """Set both ASCII and CJK fonts on a run element."""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name_en
    # CJK fallback via rPr element
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading_style(doc, text, level=1):
    """Add a heading with CJK font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, font_name_cn="黑体", font_name_en="Calibri",
                      size_pt={1: 22, 2: 16, 3: 13}.get(level, 12),
                      bold=True)
    return heading


def _add_para(doc, text, bold=False, size_pt=11, alignment=None, spacing_after=6,
              color=None, font_cn="宋体", font_en="Calibri"):
    """Add a paragraph with CJK font and optional color."""
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(spacing_after)
    run = para.add_run(text)
    _set_run_font(run, font_name_cn=font_cn, font_name_en=font_en,
                  size_pt=size_pt, bold=bold, color=color)
    return para


def _set_cell_font(cell, text, bold=False, size_pt=10):
    """Set font on a table cell."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    _set_run_font(run, font_name_cn="宋体", font_name_en="Calibri",
                  size_pt=size_pt, bold=bold)


def _set_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'CCCCCC')
        borders.append(element)
    tblPr.append(borders)


def _add_chapter_header_row(table, chapter_title, chapter_index):
    """Add a bold chapter header row spanning both columns."""
    row = table.add_row()
    # Merge the two cells into one
    cell_a = row.cells[0]
    cell_b = row.cells[1]
    cell_a.merge(cell_b)
    merged = row.cells[0]
    merged.text = ""
    para = merged.paragraphs[0]
    run = para.add_run(f"#{chapter_index} {chapter_title}")
    _set_run_font(run, font_name_cn="黑体", font_name_en="Calibri",
                  size_pt=10, bold=True, color=(251, 114, 153))
    # Light pink background for chapter header
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF0F5')
    shading.set(qn('w:val'), 'clear')
    merged._element.get_or_add_tcPr().append(shading)


def generate_docx(
    title: str,
    bvid: str,
    author: str,
    summary: str = "",
    transcript: str = "",
    tags: list = None,
    *,
    chapters: list = None,
) -> bytes:
    """
    Generate a complete DOCX document and return as bytes.

    Template structure:
      1. Cover page    — title, BV号, UP主, 导出日期
      2. AI 完整总结   — the summary text
      3. 字幕文字稿     — transcript as a table (timestamp | content)
      4. 标签          — tag list

    Args:
        title:      视频标题
        bvid:       BV号
        author:     UP主名称
        summary:    AI summary text (can be multiline)
        transcript: Full transcript text (raw, with timestamps)
        tags:       List of tag strings
        chapters:   Optional list of dicts: [{from, to, title}, ...]
                    When provided, subtitles are grouped under chapter headings
                    with a "其他片段" section for unmatched lines.

    Returns:
        bytes of the generated .docx file
    """
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =====================================================================
    # 1. Cover Page
    # =====================================================================
    # Blank space for centering
    for _ in range(4):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("学霸笔记")
    _set_run_font(run, font_name_cn="黑体", font_name_en="Calibri",
                  size_pt=32, bold=True,
                  color=(251, 114, 153))  # B站粉色

    # Divider line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("━" * 30)
    _set_run_font(run, font_name_cn="宋体", font_name_en="Calibri",
                  size_pt=10, color=(200, 200, 200))

    _add_para(doc, title, bold=True, size_pt=18,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)
    _add_para(doc, f"UP主：{author}", size_pt=12,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4)
    _add_para(doc, f"BV号：{bvid}", size_pt=11,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4)
    _add_para(doc, f"链接：https://www.bilibili.com/video/{bvid}", size_pt=10,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4)

    from datetime import datetime
    date_str = datetime.now().strftime("%Y年%m月%d日")
    _add_para(doc, f"导出日期：{date_str}", size_pt=10,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4)
    _add_para(doc, "由 BiliSum AI 生成", size_pt=9,
              alignment=WD_ALIGN_PARAGRAPH.CENTER,
              color=(150, 150, 150))

    # Page break after cover
    doc.add_page_break()

    # =====================================================================
    # 2. AI 完整总结
    # =====================================================================
    _add_heading_style(doc, "AI 完整总结", level=1)

    if summary:
        for para_text in summary.strip().split("\n"):
            trimmed = para_text.strip()
            if not trimmed:
                doc.add_paragraph()
                continue
            # Detect markdown headings and render as bold
            if trimmed.startswith("## "):
                _add_heading_style(doc, trimmed[3:].strip(), level=2)
            elif trimmed.startswith("### "):
                _add_heading_style(doc, trimmed[4:].strip(), level=3)
            elif trimmed.startswith("- "):
                _add_para(doc, "  " + trimmed, size_pt=11, spacing_after=2)
            else:
                _add_para(doc, trimmed, size_pt=11, spacing_after=4)
    else:
        _add_para(doc, "(暂无AI总结内容)", size_pt=10)

    doc.add_page_break()

    # =====================================================================
    # 3. 字幕文字稿 (table format, chapter-aware)
    # =====================================================================
    _add_heading_style(doc, "字幕文字稿", level=1)

    if transcript and transcript.strip():
        # Parse transcript: expected format with timestamps like "[MM:SS] content"
        import re as _re
        ts_pattern = _re.compile(r'^\[(\d{1,2}:\d{2}(?::\d{2})?)\]\s*(.*)')
        raw_lines = transcript.strip().split("\n")

        # Build subtitle entries from raw lines (with seconds parsed)
        subtitle_entries = []
        for line in raw_lines:
            m = ts_pattern.match(line.strip())
            if m:
                ts_str, content = m.group(1), m.group(2).strip()
                # Parse timestamp to seconds
                parts = ts_str.split(":")
                if len(parts) == 2:
                    secs = int(parts[0]) * 60 + int(parts[1])
                else:
                    secs = int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
                subtitle_entries.append({
                    "ts": ts_str,
                    "content": content,
                    "from_seconds": secs
                })
            else:
                txt = line.strip()
                if txt and len(txt) > 2:
                    subtitle_entries.append({
                        "ts": "",
                        "content": txt,
                        "from_seconds": -1
                    })

        # Helper to build flat table
        def _build_flat_table():
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(table)
            hdr_cells = table.rows[0].cells
            _set_cell_font(hdr_cells[0], "时间戳", bold=True, size_pt=10)
            _set_cell_font(hdr_cells[1], "内容", bold=True, size_pt=10)
            for cell in hdr_cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'FB7299')
                shading.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)
            row_count = 1
            for ent in subtitle_entries:
                row = table.add_row()
                _set_cell_font(row.cells[0], ent["ts"], bold=False, size_pt=9)
                _set_cell_font(row.cells[1], ent["content"], bold=False, size_pt=9)
                row_count += 1
                if row_count > 500:
                    break
            for row in table.rows:
                row.cells[0].width = Cm(2.5)
                row.cells[1].width = Cm(13.0)

        # Chapter-aware layout
        if chapters and len(chapters) > 0:
            # Normalize chapters: ensure sorted by 'from', compute effective 'to'
            norm_chapters = []
            for ch in chapters:
                ch_from = float(ch.get("from", ch.get("startTime", 0)) or 0)
                ch_to = float(ch.get("to", ch.get("endTime", 0)) or 0)
                ch_title = ch.get("title", "")
                if ch_title:
                    norm_chapters.append({"from": ch_from, "to": ch_to, "title": ch_title})
            norm_chapters.sort(key=lambda c: c["from"])

            # Build chapter table
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(table)
            # Header row
            hdr_cells = table.rows[0].cells
            _set_cell_font(hdr_cells[0], "时间戳", bold=True, size_pt=10)
            _set_cell_font(hdr_cells[1], "内容", bold=True, size_pt=10)
            for cell in hdr_cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'FB7299')
                shading.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(255, 255, 255)

            used_indexes = set()
            row_count = 1

            for idx, ch in enumerate(norm_chapters):
                ch_start = ch["from"]
                ch_title = ch["title"]
                # End boundary: next chapter's from, or this chapter's to, or Infinity
                ch_to_val = ch["to"]
                next_ch = norm_chapters[idx + 1] if idx + 1 < len(norm_chapters) else None
                if next_ch and next_ch["from"] > ch_start:
                    ch_end = next_ch["from"]
                elif ch_to_val > ch_start:
                    ch_end = ch_to_val
                else:
                    ch_end = float("inf")

                # Filter entries belonging to this chapter
                section_entries = [
                    (i, ent) for i, ent in enumerate(subtitle_entries)
                    if ent["from_seconds"] >= 0
                    and ent["from_seconds"] + 0.001 >= ch_start
                    and ent["from_seconds"] < ch_end
                ]

                if not section_entries:
                    continue

                # Chapter header row
                _add_chapter_header_row(table, ch_title, idx + 1)

                for orig_i, ent in section_entries:
                    used_indexes.add(orig_i)
                    row = table.add_row()
                    _set_cell_font(row.cells[0], ent["ts"], bold=False, size_pt=9)
                    _set_cell_font(row.cells[1], ent["content"], bold=False, size_pt=9)
                    row_count += 1
                    if row_count > 500:
                        break
                if row_count > 500:
                    break

            # Remaining / unmatched entries → "其他片段"
            remaining = [(i, ent) for i, ent in enumerate(subtitle_entries)
                         if i not in used_indexes and ent["from_seconds"] >= 0]
            if remaining:
                _add_chapter_header_row(table, "其他片段", len(norm_chapters) + 1)
                for orig_i, ent in remaining:
                    row = table.add_row()
                    _set_cell_font(row.cells[0], ent["ts"], bold=False, size_pt=9)
                    _set_cell_font(row.cells[1], ent["content"], bold=False, size_pt=9)
                    row_count += 1
                    if row_count > 500:
                        break

            # Set column widths
            for row in table.rows:
                row.cells[0].width = Cm(2.5)
                row.cells[1].width = Cm(13.0)
        else:
            # No chapters: flat table (original behavior)
            _build_flat_table()
    else:
        _add_para(doc, "(暂无字幕)", size_pt=10)

    doc.add_page_break()

    # =====================================================================
    # 4. 标签
    # =====================================================================
    _add_heading_style(doc, "标签", level=1)

    if tags and len(tags) > 0:
        tag_text = " · ".join(tags)
        _add_para(doc, tag_text, size_pt=11)
    else:
        _add_para(doc, "(无标签)", size_pt=10)

    # --- Footer note ---
    doc.add_paragraph()
    _add_para(doc, "本文档由 BiliSum AI 自动生成 | https://github.com", size_pt=8,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(180, 180, 180))

    # --- Serialize to bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
